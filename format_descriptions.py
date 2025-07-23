# -*- coding: utf-8 -*-
# ---------------------------------------------------------------------------
# Format Taxa Descriptions
# Author: Amanda Droghini
# Last Updated: 2025-07-22
# Usage: Execute in Python 3.13+.
# Description: "Format Taxa Descriptions" reads in Word documents, separates section headings from text, re-formats
# the document using Markdown, and adds Hugo front matter. The output is a Markdown file that can be converted to a
# HTML page using the static website generator Hugo.
# ---------------------------------------------------------------------------

# Import required libraries
from pathlib import Path
from docx import Document
import re
import polars as pl
import shutil

# Set root directory
drive = Path('C:/')
root_folder = 'ACCS_Work'

# Define folder structure
project_folder = drive / root_folder / 'Projects' / 'Lichen_Guide' / 'Guide Master Folder_V_7_16_25'
taxa_folder = project_folder / 'Taxa Folders'
website_folder = drive / root_folder / 'Servers_Websites' / 'akveg-lichens'
output_folder = website_folder / 'content' / 'taxa'

# Create empty lists for storing file paths and taxa names
taxon_data = []
image_files = []

# List all Word documents in taxa folder directory
## Exclude folders that do not start with a '_' (indicates in-progress taxa)
## Exclude hidden temp files
for subfolder in taxa_folder.iterdir():
    if subfolder.is_dir() and subfolder.name.startswith('_'):
        subfolder_string = str(subfolder)
        for docx_file in subfolder.rglob('*.docx'):
            if not docx_file.name.startswith('~$'):
                doc_string = str(docx_file)
                try:
                    document = Document(doc_string)
                    for para in document.paragraphs:
                        if re.match(pattern=r'Name:\s', string=para.text):
                            taxon_name = re.split(pattern=r'Name:\s(.*)', string=para.text)[1].strip()
                            cleaned_name = taxon_name.lower().replace(" ", "_")  # Replace whitespace
                            cleaned_name = re.sub(r"&_|\.", "", cleaned_name)  # Remove special characters
                            taxon_filename = cleaned_name + ".md"
                            output_filepath = output_folder / taxon_filename
                            taxon_data.append({"taxon_name": taxon_name,
                                               "cleaned_name": cleaned_name,
                                               "file_name": taxon_filename,
                                               "input_path": str(doc_string),
                                               "output_path": str(output_filepath),
                                               "folder_path": subfolder_string})
                except Exception as e:
                    print(f"Error opening or reading ", doc_string)
        for image in subfolder.rglob('*.[jpg][jpeg][png]'):
            img_string = str(image)
            image_files.append({"folder_path": subfolder_string,
                                "image_path": img_string,
                                "image_ext": image.suffix})

# Convert to polars df
image_list = pl.DataFrame(image_files)
taxon_list = pl.DataFrame(taxon_data)

# Create new image names

## Append cleaned name to image list
image_list = taxon_list.select(["folder_path", "cleaned_name"]).join(image_list, how="right", on="folder_path")  ##
# Not all folders have accompanying images

## Generate sequential numbering by taxon name
image_list = image_list.with_columns(
    (pl.int_range(0, pl.len()).over("cleaned_name") + 1).alias("taxon_id")
)

## Pad numbers under 10 with a leading zero
image_list = image_list.with_columns(
    pl.col("taxon_id")
    .cast(pl.String)
    .str.zfill(2)
    .alias("padded_taxon_id")
)

## Concatenate strings to create new clean file name
image_list = image_list.with_columns(
    pl.concat_str(
        [
            pl.col("cleaned_name"),
            pl.lit("_"),
            pl.col("padded_taxon_id"),
            pl.col("image_ext")
        ]
    ).alias("renamed_file")
)

## Ensure there are no nulls in this column
image_list.filter(pl.col("renamed_file").is_null()).shape[0]

########## IN PROGRESS
temp = image_list[0:2, :]

# Copy images to website folder
for image in temp['renamed_file']:
    new_path = website_folder / "static" / "images" / image
    src_path = Path(image_list.filter(pl.col('renamed_file') == image)['image_path'].item())
    shutil.copy(src_path, new_path)




# Process taxonomic description into Markdown

# Generate temp to test
temp = taxon_list[0:3]

for taxon in temp['taxon_name']:
    doc_path = taxon_list.filter(pl.col('taxon_name') == taxon)['input_path'].item()  # Keep as string
    md_path = Path(taxon_list.filter(pl.col('taxon_name') == taxon)['output_path'].item())

    # Open Word document
    document = Document(doc_path)

    with (open(md_path, 'w', encoding='utf-8') as md_file):
        print(f"Writing Markdown to {md_path}...")

        # Write front matter
front_matter = f"""---
title: "{taxon}"
type: docs
---


"""

        md_file.write(front_matter)

        # Write Heading 1

        md_file.write("# " + taxon + "\n\n")

        # Insert first image
        md_file.write("![" + taxon + "](/images/file_name.jpg)")

        # Format subsequent headings + paragraph text

        for para in document.paragraphs:
            split_para = re.split(pattern=r':\s(.*)', string=para.text, maxsplit=2)
            if len(split_para) > 1:
                para_heading = split_para[0].strip()
                para_text = split_para[1].strip()
                # Format heading
                formatted_heading = "## " + para_heading
                # Concatenate heading and paragraph text
                formatted_line = formatted_heading + "\n" + para_text + "\n\n"
                md_file.write(formatted_line)
            elif len(para.text) < 15:
                print(f"Skipping writing string", para.text)
            else:
                formatted_line = para.text + "\n\n"
                md_file.write(formatted_line)

